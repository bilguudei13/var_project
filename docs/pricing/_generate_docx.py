"""Convert all pricing markdown files to .docx using python-docx."""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD_DIR = os.path.dirname(__file__)
FILES = [f for f in sorted(os.listdir(MD_DIR)) if f.endswith(".md")]

def md_to_docx(md_path, docx_path):
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin   = int(2.54 * 914400 / 2.54)  # 2.54 cm
        section.right_margin  = int(2.54 * 914400 / 2.54)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    code_block = False
    code_lines = []

    def flush_code(doc, lines):
        para = doc.add_paragraph()
        run = para.add_run("\n".join(lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(6)

    for raw in lines:
        line = raw.rstrip("\n")

        # Code fence toggle
        if line.startswith("```"):
            if code_block:
                flush_code(doc, code_lines)
                code_lines = []
                code_block = False
            else:
                code_block = True
            continue

        if code_block:
            code_lines.append(line)
            continue

        # Strip trailing spaces
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Table row
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip separator rows
            if all(set(c.replace("-", "").replace(":", "").replace(" ", "")) <= {""} for c in cells):
                continue
            table = None
            # Reuse last table if exists
            if doc.tables:
                table = doc.tables[-1]
                row = table.add_row()
            else:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                row = table.rows[0]
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    # Remove inline code backticks and markdown bold
                    cell_text = re.sub(r"`([^`]+)`", r"\1", cell_text)
                    cell_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                    # Remove math $...$
                    cell_text = re.sub(r"\$([^$]+)\$", r"\1", cell_text)
                    row.cells[i].text = cell_text

        # Horizontal rule
        elif stripped == "---":
            doc.add_paragraph("─" * 60)

        # Bullet list
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            text = stripped[2:]
            text = re.sub(r"`([^`]+)`", r"\1", text)
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"\$([^$]+)\$", r"\1", text)
            p.add_run(text)

        # Math block ($$...$$) — render as italic paragraph
        elif stripped.startswith("$$"):
            math_text = stripped.strip("$").strip()
            p = doc.add_paragraph()
            run = p.add_run(math_text)
            run.italic = True
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        # Normal paragraph
        else:
            text = stripped
            # Remove inline markdown
            text = re.sub(r"`([^`]+)`", r"\1", text)
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"\$([^$]+)\$", r"\1", text)
            # Remove markdown links [text](url) -> text
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            doc.add_paragraph(text)

    doc.save(docx_path)
    print(f"  Saved: {os.path.basename(docx_path)}")


if __name__ == "__main__":
    for md_file in FILES:
        md_path   = os.path.join(MD_DIR, md_file)
        docx_path = os.path.join(MD_DIR, md_file.replace(".md", ".docx"))
        print(f"Converting {md_file} ...")
        md_to_docx(md_path, docx_path)
    print("Done.")
